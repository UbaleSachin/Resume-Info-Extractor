import os
import json
import time
import requests
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging
import PyPDF2
import docx
import pandas as pd
import re
from dotenv import load_dotenv
from openai import OpenAI
import asyncio
from collections import defaultdict, deque

# Configure logger
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

load_dotenv()

class ResumeExtractor:
    """Enhanced resume extractor with multi-API provider support and async compatibility"""
    
    def __init__(self):
        """Initialize the multi-API resume extractor."""
        # API configurations
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("OPENAI_API_KEY is not set")
        
        # Available API providers and models
        self.api_providers = {
            'openai': {
                'models': [('gpt-3.5-turbo', 4000)],  # Added gpt-4 as backup
                'api_key': self.openai_api_key,
            }
        }
        
        # Current provider and model
        self.current_provider = 'openai'
        self.current_model_index = 0
        
        # Usage tracking
        self.usage_file = 'resume_api_usage_tracking.json'
        self.usage_data = self._load_usage_data()
        
        # Rate limiting - aligned with main.py settings
        self.daily_limits = {
            'openai': 500  # RPM limit for gpt-3.5-turbo
        }

        # Define extraction prompt with better JSON structure
        self.extraction_prompt = """
You are an expert resume parser. Extract relevant information from the resume below and return it in **valid JSON** with the following structure:

{
  "personal_info": {
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "portfolio": ""
  },
  "summary": "",
  "skills": [],
  "experience": [
    {
      "title": "",
      "company": "",
      "location": "",
      "duration": "",
      "description": []
    }
  ],
  "education": [
    {
      "degree": "",
      "institution": "",
      "location": "",
      "year": "",
      "gpa": ""
    }
  ],
  "certifications": [
    {
      "name": "",
      "issuer": "",
      "date": "",
      "expiry": ""
    }
  ],
  "projects": [
    {
      "name": "",
      "description": "",
      "technologies": [],
      "duration": ""
    }
  ],
  "languages": [
    {
      "language": "",
      "proficiency": ""
    }
  ],
  "awards": []
}

**Rules:**
- Extract only data present in the resume
- Use empty strings for missing values
- Include full phone number with country code
- Add country code only if present in resume
- Normalize dates (MM/YYYY or YYYY format)
- Clean extra whitespace
- Extract both technical and soft skills
- Include partial email or phone if incomplete
- For experience description, use array of bullet points
- Return only the JSON object — no extra text or markdown

Resume text:
"""
        
        # Initialize OpenAI client
        self.openai_client = OpenAI(api_key=self.openai_api_key) if self.openai_api_key else None
        
        print(f"Initialized with provider: {self.current_provider}")
        print(f"Current model: {self.get_current_model()}")

    def _make_openai_api_call(self, text_content: str, model: str) -> Optional[str]:
        """Make API call to OpenAI for resume extraction with better error handling."""
        if not self.openai_client:
            logger.error("OpenAI client not initialized - check API key")
            return None
            
        try:
            full_prompt = self.extraction_prompt + text_content
            
            # Truncate content if too long to prevent token limit issues
            max_tokens_for_prompt = 3000 if model == 'gpt-3.5-turbo' else 7000
            if len(full_prompt) > max_tokens_for_prompt * 4:  # Rough character to token ratio
                text_content = text_content[:max_tokens_for_prompt * 4 - len(self.extraction_prompt)]
                full_prompt = self.extraction_prompt + text_content
            
            completion = self.openai_client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are a professional resume parser. Always return valid JSON without any markdown formatting or extra text."
                    },
                    {
                        "role": "user", 
                        "content": full_prompt
                    }
                ],
                max_tokens=2000,
                temperature=0.1,  # Slightly increased for better variation
            )

            # Update usage tracking
            self._update_api_usage('openai', model, 1)
            return completion.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error making OpenAI API call: {str(e)}")
            # Try with backup model if available
            if model == 'gpt-3.5-turbo' and self._can_use_backup_model():
                logger.info("Trying backup model gpt-4")
                return self._make_openai_api_call(text_content, 'gpt-4')
            return None

    def _can_use_backup_model(self) -> bool:
        """Check if backup model can be used based on usage limits."""
        today_usage = self._get_today_usage('openai', 'gpt-4')
        return today_usage < 50  # Lower limit for gpt-4

    def _load_usage_data(self) -> Dict:
        """Load API usage data from file."""
        try:
            if os.path.exists(self.usage_file):
                with open(self.usage_file, 'r') as f:
                    return json.load(f)
            else:
                return {}
        except Exception as e:
            logger.error(f"Error loading usage data: {e}")
            return {}

    def _save_usage_data(self):
        """Save API usage data to file."""
        try:
            with open(self.usage_file, 'w') as f:
                json.dump(self.usage_data, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving usage data: {e}")

    def _get_today_key(self) -> str:
        """Get today's date as a string key."""
        return datetime.now().strftime('%Y-%m-%d')

    def _update_api_usage(self, provider: str, model: str, count: int = 1):
        """Update API usage for a provider and model."""
        today = self._get_today_key()
        
        if provider not in self.usage_data:
            self.usage_data[provider] = {}
        
        if model not in self.usage_data[provider]:
            self.usage_data[provider][model] = {}
        
        if today not in self.usage_data[provider][model]:
            self.usage_data[provider][model][today] = 0
        
        self.usage_data[provider][model][today] += count
        self._save_usage_data()

    def _get_today_usage(self, provider: str, model: str = None) -> int:
        """Get today's API usage for a provider (and optionally specific model)."""
        today = self._get_today_key()
        
        if model:
            return self.usage_data.get(provider, {}).get(model, {}).get(today, 0)
        else:
            # Get total usage for provider across all models
            total = 0
            provider_data = self.usage_data.get(provider, {})
            for model_data in provider_data.values():
                total += model_data.get(today, 0)
            return total

    def get_current_model(self) -> str:
        """Get the current model being used."""
        provider_config = self.api_providers[self.current_provider]
        return provider_config['models'][self.current_model_index][0]

    def _make_api_call(self, text_content: str, provider: str, model: str) -> Optional[str]:
        """Make API call to extract resume data."""
        if provider == 'openai':
            return self._make_openai_api_call(text_content, model)
        else:
            logger.error(f"Unsupported provider: {provider}")
            return None

    def extract_from_file(self, file_path: str, filename: str, max_retries: int = 3) -> Dict[str, Any]:
        """
        Extract resume data from a file using multi-API approach with improved error handling
        """
        try:
            # Extract text from file based on type
            text_content = self._extract_text_from_file(file_path)
            
            if not text_content.strip():
                return {
                    "filename": filename,
                    "error": "No text content could be extracted from the file",
                    "success": False
                }
            
            # Check if we've hit rate limits
            if self._check_rate_limits():
                return {
                    "filename": filename,
                    "error": "Rate limit exceeded. Please try again later.",
                    "success": False
                }
                
            current_model = self.get_current_model()
            logger.info(f"Using {self.current_provider} - {current_model} for {filename}")
            
            # Make API call with retries
            for attempt in range(max_retries):
                try:
                    response_text = self._make_api_call(text_content, self.current_provider, current_model)
                    
                    if response_text:
                        # Clean the response text to extract JSON
                        json_text = self._clean_json_response(response_text)
                        extracted_data = json.loads(json_text)
                        
                        # Add metadata
                        extracted_data["filename"] = filename
                        extracted_data["success"] = True
                        extracted_data["text_length"] = len(text_content)
                        extracted_data["provider"] = self.current_provider
                        extracted_data["model"] = current_model
                        extracted_data["extraction_timestamp"] = datetime.now().isoformat()
                        
                        # Post-process and validate data
                        extracted_data = self._post_process_data(extracted_data)
                        
                        return extracted_data
                        
                except json.JSONDecodeError as e:
                    logger.error(f"JSON decode error on attempt {attempt + 1}: {str(e)}")
                    if attempt == max_retries - 1:
                        return {
                            "filename": filename,
                            "error": f"Failed to parse AI response as JSON after {max_retries} attempts: {str(e)}",
                            "raw_response": response_text[:500] if response_text else "No response",
                            "success": False
                        }
                    # Wait before retry
                    time.sleep(1 * (attempt + 1))
                    
                except Exception as e:
                    logger.error(f"Error on attempt {attempt + 1}: {str(e)}")
                    if attempt == max_retries - 1:
                        return {
                            "filename": filename,
                            "error": f"Failed to extract resume data after {max_retries} attempts: {str(e)}",
                            "success": False
                        }
                    time.sleep(1 * (attempt + 1))
            
            return {
                "filename": filename,
                "error": f"Failed to extract resume data after {max_retries} attempts.",
                "success": False
            }
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            return {
                "filename": filename,
                "error": f"Error processing file: {str(e)}",
                "success": False
            }

    def _check_rate_limits(self) -> bool:
        """Check if we've exceeded rate limits."""
        today_usage = self._get_today_usage(self.current_provider)
        daily_limit = self.daily_limits.get(self.current_provider, 1000)
        return today_usage >= daily_limit

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from various file formats with better error handling"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_docx(file_path)
            elif file_extension in ['.xls', '.xlsx']:
                return self._extract_from_excel(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension} file: {str(e)}")
            raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with better error handling"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        continue
                        
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return text.strip()

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with better error handling"""
        text = ""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        
        return text.strip()

    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file with better error handling"""
        text = ""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    # Convert dataframe to string, handling NaN values
                    sheet_text = df.fillna('').to_string(index=False)
                    text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
                except Exception as e:
                    logger.warning(f"Error reading sheet {sheet_name}: {e}")
                    continue
                    
        except Exception as e:
            raise Exception(f"Error reading Excel: {str(e)}")
        
        return text.strip()

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file with better encoding handling"""
        text = ""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if not text:
            raise Exception("Could not read TXT file with any supported encoding")
        
        return text.strip()

    def _clean_json_response(self, response_text: str) -> str:
        """Clean the AI response to extract valid JSON with better regex"""
        if not response_text:
            return "{}"
            
        # Remove markdown code blocks if present
        response_text = re.sub(r'```json\s*', '', response_text, flags=re.IGNORECASE)
        response_text = re.sub(r'```\s*$', '', response_text, flags=re.MULTILINE)
        
        # Remove any leading/trailing whitespace
        response_text = response_text.strip()
        
        # Find JSON content between curly braces (improved regex)
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(0)
            # Basic validation - ensure it starts and ends with braces
            if json_text.startswith('{') and json_text.endswith('}'):
                return json_text
        
        # If no valid JSON found, return the original text
        return response_text

    def _post_process_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Post-process and validate extracted data with enhanced cleaning"""
        # Ensure all required keys exist
        required_keys = ["personal_info", "skills", "experience", "education", "certifications", "projects", "languages", "awards"]
        for key in required_keys:
            if key not in data or data[key] is None:
                if key == "personal_info":
                    data[key] = {}
                else:
                    data[key] = []
        
        # Clean up personal info
        if "personal_info" in data and data["personal_info"]:
            personal_info = data["personal_info"]
            
            # Clean email with improved regex
            if "email" in personal_info and personal_info["email"]:
                email = personal_info["email"]
                email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', email)
                if email_match:
                    personal_info["email"] = email_match.group(0).lower()
            
            # Clean phone number with improved regex
            if "phone" in personal_info and personal_info["phone"]:
                phone = personal_info["phone"]
                # Enhanced phone number extraction
                phone_match = re.search(
                    r'(\+?\d{1,3}[\s\-\.]?)?(\(?\d{2,4}\)?[\s\-\.]?)?(\d{3,4}[\s\-\.]?\d{3,4}[\s\-\.]?\d{0,4})',
                    phone
                )
                if phone_match:
                    # Clean and format phone number
                    extracted = ''.join(filter(None, phone_match.groups()))
                    cleaned = re.sub(r'[^\d\+\-\(\)]', '', extracted)
                    personal_info["phone"] = cleaned
            
            # Clean LinkedIn URL
            if "linkedin" in personal_info and personal_info["linkedin"]:
                linkedin = personal_info["linkedin"]
                if "linkedin.com" in linkedin.lower():
                    personal_info["linkedin"] = linkedin
                elif linkedin.startswith("linkedin.com") or linkedin.startswith("www.linkedin.com"):
                    personal_info["linkedin"] = "https://" + linkedin
                else:
                    # Try to extract LinkedIn username
                    linkedin_match = re.search(r'linkedin\.com/in/([^/\s]+)', linkedin)
                    if linkedin_match:
                        personal_info["linkedin"] = f"https://linkedin.com/in/{linkedin_match.group(1)}"
        
        # Clean up skills - remove duplicates and empty strings
        if "skills" in data and isinstance(data["skills"], list):
            skills = []
            seen_skills = set()
            for skill in data["skills"]:
                if skill and skill.strip():
                    skill_cleaned = skill.strip().title()
                    if skill_cleaned.lower() not in seen_skills:
                        skills.append(skill_cleaned)
                        seen_skills.add(skill_cleaned.lower())
            data["skills"] = skills
        
        # Clean up experience descriptions
        if "experience" in data and isinstance(data["experience"], list):
            for exp in data["experience"]:
                if isinstance(exp, dict) and "description" in exp:
                    if isinstance(exp["description"], str):
                        # Convert string to list if needed
                        exp["description"] = [exp["description"]]
                    elif isinstance(exp["description"], list):
                        # Clean up list items
                        exp["description"] = [desc.strip() for desc in exp["description"] if desc and desc.strip()]
        
        # Add summary if missing but experience exists
        if not data.get("summary") and data.get("experience"):
            try:
                first_exp = data["experience"][0]
                if isinstance(first_exp, dict) and first_exp.get("title"):
                    data["summary"] = f"Professional with experience as {first_exp['title']}"
            except (IndexError, KeyError):
                pass
        
        return data

    def get_usage_summary(self) -> Dict:
        """Get usage summary for all providers and models."""
        today = self._get_today_key()
        summary = {}
        
        for provider, config in self.api_providers.items():
            if not config['api_key']:
                continue
                
            provider_summary = {}
            total_usage = self._get_today_usage(provider)
            daily_limit = self.daily_limits.get(provider, 1000)
            
            provider_summary['total_usage'] = total_usage
            provider_summary['daily_limit'] = daily_limit
            provider_summary['remaining'] = daily_limit - total_usage
            provider_summary['usage_percentage'] = (total_usage / daily_limit) * 100
            provider_summary['status'] = 'Available' if total_usage < daily_limit else 'Exhausted'
            provider_summary['models'] = {}
            
            # Model-specific usage
            for model_info in config['models']:
                model = model_info[0]
                model_usage = self._get_today_usage(provider, model)
                provider_summary['models'][model] = {
                    'usage': model_usage,
                    'last_used': 'Today' if model_usage > 0 else 'Not used today'
                }
            
            summary[provider] = provider_summary
        
        return summary

    def extract_batch(self, file_paths: List[str], output_file: str = None) -> List[Dict[str, Any]]:
        """Extract data from multiple resume files with better progress tracking"""
        results = []
        total_files = len(file_paths)
        
        for i, file_path in enumerate(file_paths):
            filename = os.path.basename(file_path)
            logger.info(f"Processing {i+1}/{total_files}: {filename}")
            
            result = self.extract_from_file(file_path, filename)
            results.append(result)
            
            # Show progress
            if result["success"]:
                logger.info(f"✓ Successfully extracted data from {filename}")
                if "personal_info" in result and result["personal_info"].get("name"):
                    logger.info(f"  Candidate: {result['personal_info']['name']}")
            else:
                logger.warning(f"✗ Failed to extract data from {filename}: {result.get('error', 'Unknown error')}")
            
            # Small delay to avoid overwhelming the API
            time.sleep(0.1)
        
        # Save results if output file specified
        if output_file:
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                logger.info(f"Results saved to: {output_file}")
            except Exception as e:
                logger.error(f"Error saving results: {e}")
        
        # Print summary
        successful = sum(1 for r in results if r["success"])
        logger.info(f"Batch processing complete: {successful}/{total_files} files processed successfully")
        
        return results

    async def extract_batch_async(self, file_paths: List[str], progress_callback=None) -> List[Dict[str, Any]]:
        """Async version of batch extraction for better integration with FastAPI"""
        results = []
        total_files = len(file_paths)
        
        for i, file_path in enumerate(file_paths):
            filename = os.path.basename(file_path)
            
            # Run extraction in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, self.extract_from_file, file_path, filename)
            results.append(result)
            
            # Call progress callback if provided
            if progress_callback:
                progress = int((i + 1) * 100 / total_files)
                await progress_callback(progress)
            
            # Small delay to avoid overwhelming the API
            await asyncio.sleep(0.1)
        
        return results