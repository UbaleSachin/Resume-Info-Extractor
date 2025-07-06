import os
import json
import time
import requests
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging
import PyPDF2
import docx
import pandas as pd
import re
from dotenv import load_dotenv
from openai import OpenAI

# Configure logger
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

load_dotenv()

import time
from collections import defaultdict, deque

class ModelRateLimiter:
    def __init__(self, rpm_limits):
        self.rpm_limits = rpm_limits  # {'model_name': rpm}
        self.request_times = defaultdict(deque)  # {'model_name': deque([timestamps])}

    def can_send(self, model_name):
        now = time.time()
        window = 60  # seconds
        q = self.request_times[model_name]
        # Remove requests older than 60 seconds
        while q and q[0] < now - window:
            q.popleft()
        return len(q) < self.rpm_limits[model_name]

    def record(self, model_name):
        self.request_times[model_name].append(time.time())

    def wait_for_slot(self, model_name):
        while not self.can_send(model_name):
            time.sleep(0.5)

rate_limiter = ModelRateLimiter({
    'llama3-70b-8192': 30,
    'llama3-8b-8192': 30,
    'llama-3.3-70b-versatile': 30,
    'mixtral-saba-24b': 30,
    'gemma-3-27b-it': 30,
    'gemma-3-12b-it': 30,
    'gemma-3n-e4b-it': 30,
    'gemma-3n-e2b-it': 30,
    'gemini-2.0-flash-lite': 30,
    'gemini-2.0-flash': 15,
    'gemini-2.5-flash-preview-04-17': 10,
    'gpt-3.5-turbo': 500,
})

class ResumeExtractor:
    """Enhanced resume extractor with multi-API provider support"""
    
    def __init__(self):
        """Initialize the multi-API resume extractor."""
        
        # API configurations
        self.groq_api_key = os.getenv('GROQ_API_KEY')
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        self.gemini_api_key = os.getenv('GEMINI_API_KEY')
        
        # Available API providers and models
        self.api_providers = {
            'groq': {
                'base_url': 'https://api.groq.com/openai/v1',
                'models': [
                    ('llama3-70b-8192', 14400),
                    ('llama3-8b-8192', 14400),
                    ('llama-3.3-70b-versatile', 1000),
                    ('mixtral-saba-24b', 1000),
                ],
                'api_key': self.groq_api_key
            },
            'gemini': {
                'base_url': 'https://generativelanguage.googleapis.com/v1beta',
                'models': [
                    ('gemma-3-27b-it', 14400),
                    ('gemma-3-12b-it', 14400),
                    ('gemma-3n-e4b-it', 14400),
                    ('gemma-3n-e2b-it', 14400),
                    ('gemini-2.0-flash-lite', 1500),
                    ('gemini-2.0-flash', 1500),
                    ('gemini-2.5-flash-preview-04-17', 500),
                ],
                'api_key': self.gemini_api_key
            },
            'openai': {
                'models': [('gpt-3.5-turbo', 10000),],
                'api_key': self.openai_api_key,
            }
        }
        
        # Current provider and model
        self.current_provider = 'openai'  # Default to Gemini
        self.current_model_index = 0
        
        # Determine which provider to use based on available API keys
        if not self.gemini_api_key and self.groq_api_key:
            self.current_provider = 'groq'
        elif not self.gemini_api_key and not self.groq_api_key and self.openai_api_key:
            self.current_provider = 'openai'
        elif not any([self.gemini_api_key, self.groq_api_key, self.openai_api_key]):
            raise ValueError("At least one API key must be set: GEMINI_API_KEY, GROQ_API_KEY, or OPENAI_API_KEY")
        
        # Usage tracking
        self.usage_file = 'resume_api_usage_tracking.json'
        self.usage_data = self._load_usage_data()

        
        # Define extraction prompt
        self.extraction_prompt = """
            You are an expert resume parser. Extract relevant information from the resume below and return it in **valid JSON** with the following structure:

            {
            "personal_info": {
                "name": "",
                "email": "",
                "phone": "",
                "location": "",
                "linkedin": "",
                "portfolio": ""
            },
            "summary": "",
            "skills": [],
            "experience": [
                {
                "title": "",
                "company": "",
                "location": "",
                "duration": "",
                "description": ["", "", ""]
                }
            ],
            "education": [
                {
                "degree": "",
                "institution": "",
                "location": "",
                "year": "",
                "gpa": ""
                }
            ],
            "certifications": [
                {
                "name": "",
                "issuer": "",
                "date": "",
                "expiry": ""
                }
            ],
            "projects": [
                {
                "name": "",
                "description": "",
                "technologies": [],
                "duration": ""
                }
            ],
            "languages": [
                {
                "language": "",
                "proficiency": ""
                }
            ],
            "awards": []
            }

            **Rules:**
            - Extract only data present in the resume
            - Use empty strings for missing values
            - Include full phone number with country code (e.g., +1, +91) if present
            - Normalize dates (MM/YYYY or YYYY)
            - Clean extra whitespace
            - Extract both technical and soft skills
            - Include partial email or phone if incomplete
            - Return only the JSON object — no extra text

            Resume text:
            """

        
        print(f"Initialized with provider: {self.current_provider}")
        print(f"Current model: {self.get_current_model()}")

    def _make_openai_api_call(self, text_content: str, model: str) -> Optional[str]:
        """Make API call to openAI for resume extraction."""
        try:
            full_prompt = self.extraction_prompt + text_content
            client = OpenAI(api_key = self.openai_api_key)
            completion = client.chat.completions.create(
                model = model,
                messages = [
                    {
                        "role": "user", "content": full_prompt,
                    }
                ],
                max_tokens = 2000,
                temperature = 0  # Lower temperature for more consistent JSON output
            )

            # upadte usage tracking
            self._update_api_usage('openai', model, 1)
            return completion.choices[0].message.content
        except Exception as e:
            logger.error(f"Error making OpenAI API call: {str(e)}")
            return None

    def _load_usage_data(self) -> Dict:
        """Load API usage data from file."""
        try:
            if os.path.exists(self.usage_file):
                with open(self.usage_file, 'r') as f:
                    return json.load(f)
            else:
                return {}
        except Exception as e:
            print(f"Error loading usage data: {e}")
            return {}

    def _save_usage_data(self):
        """Save API usage data to file."""
        try:
            with open(self.usage_file, 'w') as f:
                json.dump(self.usage_data, f, indent=2)
        except Exception as e:
            print(f"Error saving usage data: {e}")

    def _get_today_key(self) -> str:
        """Get today's date as a string key."""
        return datetime.now().strftime('%Y-%m-%d')

    def _update_api_usage(self, provider: str, model: str, count: int = 1):
        """Update API usage for a provider and model."""
        today = self._get_today_key()
        
        if provider not in self.usage_data:
            self.usage_data[provider] = {}
        
        if model not in self.usage_data[provider]:
            self.usage_data[provider][model] = {}
        
        if today not in self.usage_data[provider][model]:
            self.usage_data[provider][model][today] = 0
        
        self.usage_data[provider][model][today] += count
        self._save_usage_data()

    def _get_today_usage(self, provider: str, model: str = None) -> int:
        """Get today's API usage for a provider (and optionally specific model)."""
        today = self._get_today_key()
        
        if model:
            return self.usage_data.get(provider, {}).get(model, {}).get(today, 0)
        else:
            # Get total usage for provider across all models
            total = 0
            provider_data = self.usage_data.get(provider, {})
            for model_data in provider_data.values():
                total += model_data.get(today, 0)
            return total

    def _can_use_provider(self, provider: str, count: int = 1) -> bool:
        """Check if a provider can be used without exceeding daily limit."""
        # Check if current model for provider is within daily limit
        current_model, daily_limit = self.api_providers[provider]['models'][self.current_model_index]
        current_usage = self._get_today_usage(provider, current_model)
        return (current_usage + count) <= daily_limit

    def get_current_model(self) -> str:
        """Get the current model being used."""
        provider_config = self.api_providers[self.current_provider]
        return provider_config['models'][self.current_model_index][0]

    def _make_gemini_api_call(self, text_content: str, model: str) -> Optional[str]:
        """Make API call to Gemini for resume extraction."""
        try:
            full_prompt = self.extraction_prompt + text_content
            
            # Prepare request payload for Gemini
            url = f"{self.api_providers['gemini']['base_url']}/models/{model}:generateContent"
            params = {'key': self.gemini_api_key}
            
            payload = {
                "contents": [
                    {
                        "parts": [
                            {
                                "text": full_prompt
                            }
                        ]
                    }
                ]
            }
            
            headers = {
                'Content-Type': 'application/json'
            }
            
            # Make API request
            response = requests.post(url, params=params, headers=headers, json=payload, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and len(result['candidates']) > 0:
                    content = result['candidates'][0]['content']
                    if 'parts' in content and len(content['parts']) > 0:
                        extracted_text = content['parts'][0]['text']
                        
                        # Update usage tracking
                        self._update_api_usage('gemini', model, 1)
                        
                        return extracted_text
                else:
                    logger.error(f"Unexpected Gemini response format: {result}")
                    return None
            else:
                logger.error(f"Gemini API call failed: {response.status_code} - {response.text}")
                return None
                
        except Exception as e:
            logger.error(f"Error making Gemini API call: {str(e)}")
            return None

    def _make_api_call(self, text_content: str, provider: str, model: str) -> Optional[str]:
        """Make API call to extract resume data."""
        if provider == 'gemini':
            return self._make_gemini_api_call(text_content, model)
        elif provider == 'openai':
            return self._make_openai_api_call(text_content, model)
        
        try:
            provider_config = self.api_providers[provider]
            full_prompt = self.extraction_prompt + text_content
            
            # Prepare request payload
            headers = {
                'Authorization': f'Bearer {provider_config["api_key"]}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                'model': model,
                'messages': [
                    {
                        'role': 'user',
                        'content': full_prompt
                    }
                ],
                'max_tokens': 2000,
                'temperature': 0  # Lower temperature for more consistent JSON output
            }
            
            # Make API request
            url = f"{provider_config['base_url']}/chat/completions"
            response = requests.post(url, headers=headers, json=payload, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                extracted_text = result['choices'][0]['message']['content']
                
                # Update usage tracking
                self._update_api_usage(provider, model, 1)
                
                return extracted_text
            else:
                logger.error(f"API call failed: {response.status_code} - {response.text}")
                return None
                
        except Exception as e:
            logger.error(f"Error making API call to {provider}: {str(e)}")
            return None

    def _switch_to_next_model(self) -> bool:
        """Switch to next available model or provider."""
        # Try next model in current provider
        provider_config = self.api_providers[self.current_provider]
        if self.current_model_index < len(provider_config['models']) - 1:
            self.current_model_index += 1
            return True
        
        # Try switching provider based on priority: gemini -> groq -> openai
        available_providers = []
        if self.gemini_api_key and self.current_provider != 'gemini':
            available_providers.append('gemini')
        if self.groq_api_key and self.current_provider != 'groq':
            available_providers.append('groq')
        if self.openai_api_key and self.current_provider != 'openai':
            available_providers.append('openai')
        
        if available_providers:
            self.current_provider = available_providers[0]
            self.current_model_index = 0
            return True
        
        return False

    def extract_from_file(self, file_path: str, filename: str, max_retries: int = 3) -> Dict[str, Any]:
        """
        Extract resume data from a file using multi-API approach
        """
        try:
            # Extract text from file based on type
            text_content = self._extract_text_from_file(file_path)
            
            if not text_content.strip():
                return {
                    "filename": filename,
                    "error": "No text content could be extracted from the file",
                    "success": False
                }
            
            # Try multiple providers/models with fallback
            for attempt in range(max_retries):
                # Check if current provider has capacity
                if not self._can_use_provider(self.current_provider, 1):
                    print(f"Provider {self.current_provider} would exceed daily limit. Switching...")
                    if not self._switch_to_next_model():
                        return {
                            "filename": filename,
                            "error": "All API providers have reached their daily limits.",
                            "success": False
                        }
                
                current_model = self.get_current_model()
                print(f"Using {self.current_provider} - {current_model} (Attempt {attempt + 1})")
                
                # Make API call
                rate_limiter.wait_for_slot(current_model)  # Ensure we don't exceed model-specific RPM
                response_text = self._make_api_call(text_content, self.current_provider, current_model)
                rate_limiter.record(current_model)  # Record the request time
                
                if response_text:
                    try:
                        # Clean the response text to extract JSON
                        json_text = self._clean_json_response(response_text)
                        extracted_data = json.loads(json_text)
                        
                        # Add metadata
                        extracted_data["filename"] = filename
                        extracted_data["success"] = True
                        extracted_data["text_length"] = len(text_content)
                        extracted_data["provider"] = self.current_provider
                        extracted_data["model"] = current_model
                        
                        # Post-process and validate data
                        extracted_data = self._post_process_data(extracted_data)
                        
                        return extracted_data
                        
                    except json.JSONDecodeError as e:
                        print(f"JSON decode error on attempt {attempt + 1}: {str(e)}")
                        if attempt < max_retries - 1:
                            # Try next model/provider
                            if not self._switch_to_next_model():
                                break
                        else:
                            return {
                                "filename": filename,
                                "error": f"Failed to parse AI response as JSON: {str(e)}",
                                "raw_response": response_text[:500],
                                "success": False
                            }
                else:
                    # Try next model/provider
                    if attempt < max_retries - 1:
                        if not self._switch_to_next_model():
                            break
            
            return {
                "filename": filename,
                "error": "Failed to extract resume data after multiple attempts.",
                "success": False
            }
                
        except Exception as e:
            return {
                "filename": filename,
                "error": f"Error processing file: {str(e)}",
                "success": False
            }

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from various file formats"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_docx(file_path)
            elif file_extension in ['.xls', '.xlsx']:
                return self._extract_from_excel(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return text

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        
        return text

    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                # Convert dataframe to string, handling NaN values
                sheet_text = df.fillna('').to_string(index=False)
                text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
                
        except Exception as e:
            raise Exception(f"Error reading Excel: {str(e)}")
        
        return text

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except Exception as e:
                raise Exception(f"Error reading TXT file with different encodings: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
        
        return text

    def _clean_json_response(self, response_text: str) -> str:
        """Clean the AI response to extract valid JSON"""
        # Remove markdown code blocks if present
        response_text = re.sub(r'```json\s*', '', response_text)
        response_text = re.sub(r'```\s*$', '', response_text)
        
        # Find JSON content between curly braces
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            return json_match.group(0)
        
        return response_text

    def _post_process_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Post-process and validate extracted data"""
        # Clean up personal info
        if "personal_info" in data and data["personal_info"]:
            personal_info = data["personal_info"]
            
            # Clean email
            if "email" in personal_info and personal_info["email"]:
                email = personal_info["email"]
                # Extract email using regex if it's embedded in text
                email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', email)
                if email_match:
                    personal_info["email"] = email_match.group(0)
            
            # Clean phone number
            if "phone" in personal_info and personal_info["phone"]:
                phone = personal_info["phone"]
                # Improved regex for international phone numbers
                phone_match = re.search(
                    r'(\+?\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?)?(\d{3,4}[\s\-]?\d{3,4}[\s\-]?\d{0,4})',
                    phone
                )
                if phone_match:
                    # Remove extra spaces/dashes and join groups
                    extracted = ''.join(filter(None, phone_match.groups()))
                    # Clean up: remove double spaces/dashes, keep only numbers, +, -, ()
                    cleaned = re.sub(r'[^\d\+\-\(\)]', '', extracted)
                    personal_info["phone"] = cleaned
        
        # Clean up skills - remove duplicates and empty strings
        if "skills" in data and isinstance(data["skills"], list):
            skills = []
            for skill in data["skills"]:
                if skill and skill.strip() and skill.strip() not in skills:
                    skills.append(skill.strip().title())  # Normalize case
            data["skills"] = skills
        
        # Ensure arrays exist even if empty
        for key in ["skills", "experience", "education", "certifications", "projects", "languages", "awards"]:
            if key not in data or data[key] is None:
                data[key] = []
        
        # Ensure personal_info exists
        if "personal_info" not in data or data["personal_info"] is None:
            data["personal_info"] = {}
        
        return data

    def get_usage_summary(self) -> Dict:
        """Get usage summary for all providers and models."""
        today = self._get_today_key()
        summary = {}
        
        for provider, config in self.api_providers.items():
            if not config['api_key']:
                continue
                
            provider_summary = {}
            total_usage = self._get_today_usage(provider)
            daily_limit = self.daily_limits.get(provider, 1000)
            
            provider_summary['total_usage'] = total_usage
            provider_summary['daily_limit'] = daily_limit
            provider_summary['remaining'] = daily_limit - total_usage
            provider_summary['usage_percentage'] = (total_usage / daily_limit) * 100
            provider_summary['status'] = 'Available' if total_usage < daily_limit else 'Exhausted'
            provider_summary['models'] = {}
            
            # Model-specific usage
            for model in config['models']:
                model_usage = self._get_today_usage(provider, model)
                provider_summary['models'][model] = {
                    'usage': model_usage,
                    'last_used': 'Today' if model_usage > 0 else 'Not used today'
                }
            
            summary[provider] = provider_summary
        
        return summary

    def switch_provider(self, provider: str, model_index: int = 0) -> bool:
        """Manually switch to a specific provider and model."""
        if provider in self.api_providers and self.api_providers[provider]['api_key']:
            self.current_provider = provider
            self.current_model_index = min(model_index, len(self.api_providers[provider]['models']) - 1)
            print(f"Switched to {provider} - {self.get_current_model()}")
            return True
        else:
            print(f"Provider {provider} not available or API key not set")
            return False

    def list_available_providers(self) -> Dict:
        """List all available providers and their models."""
        available = {}
        for provider, config in self.api_providers.items():
            if config['api_key']:
                available[provider] = {
                    'models': config['models'],
                    'status': 'Available' if self._can_use_provider(provider) else 'Daily limit reached'
                }
        return available

    def extract_batch(self, file_paths: List[str], output_file: str = None) -> List[Dict[str, Any]]:
        """Extract data from multiple resume files"""
        results = []
        
        for i, file_path in enumerate(file_paths):
            filename = os.path.basename(file_path)
            print(f"\nProcessing {i+1}/{len(file_paths)}: {filename}")
            
            result = self.extract_from_file(file_path, filename)
            results.append(result)
            
            # Show progress
            if result["success"]:
                print(f"✓ Successfully extracted data from {filename}")
                if "personal_info" in result and result["personal_info"].get("name"):
                    print(f"  Candidate: {result['personal_info']['name']}")
            else:
                print(f"✗ Failed to extract data from {filename}: {result.get('error', 'Unknown error')}")
            
            # Small delay to avoid rate limiting
            time.sleep(0.5)
        
        # Save results if output file specified
        if output_file:
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                print(f"\nResults saved to: {output_file}")
            except Exception as e:
                print(f"Error saving results: {e}")
        
        # Print summary
        successful = sum(1 for r in results if r["success"])
        print(f"\nBatch processing complete: {successful}/{len(file_paths)} files processed successfully")
        
        return results


# Example usage
if __name__ == "__main__":
    try:
        extractor = ResumeExtractor()
        
        # Show available providers
        print("Available Providers:")
        providers = extractor.list_available_providers()
        for provider, info in providers.items():
            print(f"\n{provider.upper()}:")
            print(f"  Status: {info['status']}")
            print(f"  Models: {', '.join(info['models'])}")
        
        print("\n" + "="*50 + "\n")
        
        # Show usage summary
        print("Usage Summary:")
        usage = extractor.get_usage_summary()
        for provider, stats in usage.items():
            print(f"\n{provider.upper()}:")
            print(f"  Total Usage: {stats['total_usage']}/{stats['daily_limit']} ({stats['usage_percentage']:.1f}%)")
            print(f"  Status: {stats['status']}")
        
        print("\n" + "="*50 + "\n")
        
        # Test with a single resume file
        test_resume_path = "path_to_your_resume.pdf"  # Replace with actual resume file path
        
        if os.path.exists(test_resume_path):
            print("Extracting resume data...")
            result = extractor.extract_from_file(test_resume_path, os.path.basename(test_resume_path))
            
            if result["success"]:
                print(f"Provider used: {result.get('provider', 'Unknown')}")
                print(f"Model used: {result.get('model', 'Unknown')}")
                print("\nExtracted Resume Data:")
                print(json.dumps(result, indent=2, ensure_ascii=False))
            else:
                print(f"Error: {result.get('error', 'Unknown error')}")
        else:
            print(f"Test resume not found: {test_resume_path}")
            print("Please update the test_resume_path variable with a valid resume file path.")
            
        # Example of batch processing
        # resume_files = ["resume1.pdf", "resume2.docx", "resume3.txt"]
        # results = extractor.extract_batch(resume_files, "extracted_resumes.json")
            
    except Exception as e:
        print(f"Error initializing extractor: {str(e)}")
        print("Make sure to set at least one API key in your .env file:")
        print("- GEMINI_API_KEY")
        print("- GROQ_API_KEY")
        print("- OPENAI_API_KEY")